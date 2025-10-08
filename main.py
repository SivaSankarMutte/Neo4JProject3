import os
import time
import streamlit as st
import pandas as pd
from dotenv import load_dotenv
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.graphs import Neo4jGraph
from langchain.chains import GraphCypherQAChain
from langchain_core.runnables import RunnablePassthrough
from typing import Dict, Any

# For PDFs, Word, Images (Keep for multi-input, but extraction logic simplified)
import docx
import pytesseract
from PIL import Image
import fitz 

# --- CONFIGURATION & ENV SETUP ---
load_dotenv()

# General Config
EXCEL_FILE = "./data/tickets.xlsx"
VECTOR_STORE_DIR = "vector_store"
LAST_MOD_TIME_FILE = "last_mod_time.txt"

# Environment Variables (Local)
os.environ['GROQ_API_KEY'] = os.getenv("GROQ_API_KEY")
os.environ['HF_TOKEN'] = os.getenv("HF_TOKEN")
# Neo4j Config
NEO4J_URI = os.getenv("NEO4J_URI")
NEO4J_USERNAME = os.getenv("NEO4J_USERNAME")
NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD")
NEO4J_DATABASE = os.getenv("NEO4J_DATABASE", "neo4j") # default

# --- INITIALIZATION ---

# HuggingFace embeddings
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2",
    model_kwargs={"device": "cpu"}
)

# LLM
llm = ChatGroq(
    groq_api_key=os.environ["GROQ_API_KEY"],
    model_name="llama-3.1-8b-instant",
    temperature=0
)

# Neo4j Graph
neo4j_graph = None
try:
    if NEO4J_URI and NEO4J_USERNAME and NEO4J_PASSWORD:
        neo4j_graph = Neo4jGraph(
            url=NEO4J_URI,
            username=NEO4J_USERNAME,
            password=NEO4J_PASSWORD,
            database=NEO4J_DATABASE
        )
        st.sidebar.success("🌐 Connected to Neo4j.")
    else:
        st.sidebar.error("❌ Neo4j environment variables are missing.")
except Exception as e:
    st.sidebar.error(f"❌ Failed to connect to Neo4j: {e}")

# Prompt template
prompt = ChatPromptTemplate.from_template("""
You are an expert IT support assistant. 
Use the Ticket Data (Context) and the IT Knowledge Graph Results (Graph Context) to provide a comprehensive answer.

Ticket Data (Context):
{context}

IT Knowledge Graph Results (Graph Context):
{graph_context}

Question: {input}

If there's no helpful information in the context, reply with:
"Not enough historical data or graph relationships to answer accurately."
""")


# --- UTILITY FUNCTIONS ---

def get_last_mod_time():
    """Gets the last modification time of the Excel file."""
    return os.path.getmtime(EXCEL_FILE)

def read_last_saved_mod_time():
    """Reads the last saved modification time from the tracker file."""
    if os.path.exists(LAST_MOD_TIME_FILE):
        with open(LAST_MOD_TIME_FILE, "r") as f:
            try:
                return float(f.read().strip())
            except ValueError:
                return None
    return None

def save_last_mod_time(mod_time):
    """Saves the current modification time."""
    with open(LAST_MOD_TIME_FILE, "w") as f:
        f.write(str(mod_time))
        
def load_vector_store():
    """Loads the FAISS vector store."""
    return FAISS.load_local(VECTOR_STORE_DIR, embeddings, allow_dangerous_deserialization=True)

def extract_text_from_file(file):
    """Placeholder for file extraction logic (PDF/DOCX/Image)."""
    if not file: return ""
    # In a real app, this logic would extract text from various file types
    try:
        # Simple fallback for text files
        file.seek(0) # Rewind file pointer
        return file.read().decode("utf-8")
    except Exception:
        # Placeholder logic for other file types
        return "Content from uploaded file." 


def build_vector_store_and_graph():
    """
    Builds the FAISS vector store AND populates the Neo4j Graph.
    Includes FIX: Ensures all data is converted to string to prevent embedding errors.
    """
    df = pd.read_excel(EXCEL_FILE)

    # CHECK FOR REQUIRED COLUMNS
    required_cols = ["Ticket ID", "Description", "Category", "Subcategory", "Application", "Resolution"]
    for col in required_cols:
        if col not in df.columns:
            st.error(f"❌ '{col}' column not found in Excel file. Please ensure it exists.")
            st.stop()
    
    # Drop rows where 'Description' is missing
    df.dropna(subset=["Description"], inplace=True)

    # 1. Build Vector Store (FAISS)
    st.info("Building FAISS Vector DB...")
    documents = []
    graph_nodes = []
    
    for _, row in df.iterrows():
        # --- FIX: EXPLICITLY CONVERT ALL DATA TO STRING ---
        ticket_id = str(row.get('Ticket ID', ''))
        description = str(row.get('Description', ''))
        category = str(row.get('Category', ''))
        subcategory = str(row.get('Subcategory', ''))
        application = str(row.get('Application', ''))
        resolution = str(row.get('Resolution', ''))
        
        # Document content for FAISS (RAG) - All components are strings
        content = f"""
        Ticket ID: {ticket_id}
        Description: {description}
        Category: {category}
        Subcategory: {subcategory}
        Application: {application} 
        Resolution: {resolution}
        """
        documents.append(Document(page_content=content.strip()))
        
        # Data for Neo4j (Graph) - All components are strings
        graph_nodes.append({
            'id': ticket_id,
            'description': description,
            'category': category or 'Unknown',
            'subcategory': subcategory or 'Unknown',
            'application': application or 'Unknown', 
            'resolution': resolution
        })

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    split_docs = text_splitter.split_documents(documents)
    
    # Pass split_docs (containing only string content) to FAISS
    vector_store = FAISS.from_documents(split_docs, embeddings)
    vector_store.save_local(VECTOR_STORE_DIR)
    st.success("✅ Vector DB rebuilt.")

    # 2. Populate Neo4j Graph
    if neo4j_graph:
        st.info("Populating Neo4j Knowledge Graph...")
        neo4j_graph.query("MATCH (n) DETACH DELETE n") 

        for node in graph_nodes:
            # Basic sanitization for Cypher strings (safe since they are strings)
            desc_safe = node['description'].replace("'", "\\'") 
            res_safe = node['resolution'].replace("'", "\\'")

            neo4j_graph.query(f"""
            MERGE (c:Category {{name: '{node['category']}'}})
            MERGE (s:Subcategory {{name: '{node['subcategory']}'}})
            MERGE (a:Application {{name: '{node['application']}'}})
            
            MERGE (t:Ticket {{
                ticketId: '{node['id']}', 
                description: '{desc_safe}', 
                resolution: '{res_safe}'
            }})
            
            MERGE (t)-[:BELONGS_TO_CATEGORY]->(c)
            MERGE (t)-[:BELONGS_TO_SUBCATEGORY]->(s)
            MERGE (t)-[:AFFECTS_APPLICATION]->(a) 
            """)

        st.success("✅ Neo4j Graph populated.")

# --- DYNAMIC GRAPH RETRIEVAL CHAIN (COMPATIBILITY FIX) ---

def get_dynamic_graph_context(user_query: str) -> Document:
    """Uses LLM to generate and execute a Cypher query against the Neo4j graph."""
    if not neo4j_graph:
        return Document(page_content="Neo4j not available.")

    # 1. FORCE THE GRAPH SCHEMA TO BE REFRESHED
    # This is a good practice to ensure the LLM has the latest schema.
    try:
        neo4j_graph.refresh_schema() 
    except Exception as e:
        st.warning(f"Failed to refresh Neo4j schema: {e}")
        
    # --- Custom Cypher Generation Prompt Template ---
    # This template is structured to tell the LLM to use WHERE/CONTAINS
    # It must include {schema} and {question} placeholders for the chain to work.
    custom_cypher_template = """
    You are an expert Cypher query generator for a helpdesk system. 
    Based on the provided Neo4j graph schema, write a Cypher query that precisely addresses the user's question.

    RULES:
    1. For searching ticket descriptions, use the `CONTAINS` operator or a Full-Text Search index (if available) instead of exact matching.
    2. Only return the ticket details (description, resolution, ticketId, category, application name).
    3. The schema is: {schema}

    Question: {question}
    Cypher Query:
    """
    
    # Create the LLM chain that generates the Cypher query using the custom prompt
    cypher_generator_prompt = ChatPromptTemplate.from_template(custom_cypher_template)
    
    # Instantiate the GraphCypherQAChain with the custom prompt for the Cypher step
    cypher_chain = GraphCypherQAChain.from_llm(
        llm=llm, 
        graph=neo4j_graph, 
        verbose=True, 
        allow_dangerous_requests=True,
        # This parameter controls the prompt for the Cypher generation step
        cypher_prompt=cypher_generator_prompt 
    )
    
    try:
        # Run the chain
        graph_result = cypher_chain.run(user_query)
        
        return Document(
            page_content=f"Knowledge Graph Query Result: {graph_result}",
            metadata={"source": "Neo4j Graph"}
        )
    except Exception as e:
        st.error(f"❌ Neo4j/Cypher Chain Execution Error: {e}")
        return Document(page_content="Graph query failed or provided no results due to an execution error.")

# --- COMBINED RAG CHAIN ---

def combined_retriever(query: str) -> Dict[str, Any]:
    """Fetches documents from FAISS and dynamic context from Neo4j."""
    
    # 1. Vector Store (FAISS) Retrieval
    faiss_docs = vector_store.as_retriever().invoke(query)
    
    # 2. Dynamic Graph Retrieval (returns a single Document)
    graph_context_doc = get_dynamic_graph_context(query)
    
    # Return a dictionary matching the variables expected by the main RAG prompt
    return {
        "context": faiss_docs,
        "graph_context": graph_context_doc
    }


# --- STREAMLIT UI ---

st.set_page_config(page_title="Ticket Gen AI Assistant", layout="wide")
st.title("🎫 Gen AI Ticket Assistant (FAISS + Dynamic Neo4j RAG)")

# Auto-rebuild logic
if os.path.exists(EXCEL_FILE):
    last_mod_time = get_last_mod_time()
    saved_mod_time = read_last_saved_mod_time()

    if saved_mod_time != last_mod_time or not os.path.exists(VECTOR_STORE_DIR):
        st.warning("Excel file changed — rebuilding vector DB AND Neo4j Graph...")
        build_vector_store_and_graph()
        save_last_mod_time(last_mod_time)
        # st.experimental_rerun() # Rerun removed to simplify environment setup
else:
    st.error(f"❌ Required Excel file not found: {EXCEL_FILE}. Please create the 'data' folder and {EXCEL_FILE}.")
    st.stop()


# Load existing vector store
vector_store = load_vector_store()

# Query input options
query_type = st.radio("Choose input type:", ["Text", "File (PDF/DOCX/Image)"])

user_query = ""
if query_type == "Text":
    user_query = st.text_input("📝 Describe a new ticket or ask a question:")
else:
    uploaded_query_file = st.file_uploader("📂 Upload a query file", type=["txt", "pdf", "docx", "png", "jpg", "jpeg"])
    user_query = extract_text_from_file(uploaded_query_file) if uploaded_query_file else ""

# Run query
if user_query:
    
    # 1. Define the main document chain (LLM + Prompt)
    document_chain = create_stuff_documents_chain(llm, prompt)

    # 2. Assemble the full RAG chain
    # Note: We must invoke the combined_retriever once and pass the results
    # to avoid re-running the heavy retrieval step (including the Cypher QA Chain)
    
    context_data = combined_retriever(user_query)

    retrieval_chain = (
        RunnablePassthrough.assign(
            context=lambda x: context_data['context'],
            graph_context=lambda x: context_data['graph_context'],
            input=lambda x: x['input']
        )
        | document_chain
    )
    
    try:
        with st.spinner("Thinking..."):
            # Pass only the original query here. The context is pre-fetched above.
            response = retrieval_chain.invoke({"input": user_query})
        
        st.write("### ✅ Answer")
        st.write(response)

        st.write("---")
        with st.expander("🔎 Raw Context Used"):
            st.markdown("#### FAISS Documents (Vector Similarity)")
            for doc in context_data['context']:
                st.markdown(doc.page_content)
                st.write("---")
            
            st.markdown("#### Neo4j Graph Result (Cypher QA Chain)")
            st.code(context_data['graph_context'].page_content, language="markdown")
            
    except Exception as e:
         st.error(f"An error occurred during query execution: {e}")